#!/usr/bin/env python3
from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import datetime, timezone, timedelta
import re
import subprocess
import sys
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import CondPageBreak, KeepTogether, Paragraph, SimpleDocTemplate, Spacer


ROOT = Path(__file__).resolve().parents[1]
DEFAULT_FONT = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"
REFERENCE_DOCX = Path("/Users/hannah/Downloads/科技播客情报分析报告.docx")
DOCX_FONT = "Google Sans"
DOCX_BOLD_FONT = "PingFang SC Semibold"
METADATA_LABELS = ("原始标题", "来源与发布者", "原始链接")
CONTENT_BLOCK_LABELS = ("核心内容摘要", "情报价值点", "关键金句", "证据锚点")
COMPACT_LABELS = METADATA_LABELS + CONTENT_BLOCK_LABELS
BODY_LINE_SPACING_PT = 14


@dataclass
class Section:
    title: str
    lines: list[str]


@dataclass
class TextRun:
    text: str
    bold: bool = False
    italic: bool = False


def short_date(run_id: str) -> str:
    return datetime.strptime(run_id[:8], "%Y%m%d").strftime("%y%m%d")


def scheduled_report_date(run_id: str) -> str | None:
    import json

    package_dir = ROOT / "data" / "gemini_inputs" / run_id
    source_manifest = package_dir / "source-manifest-original.json"
    manifest_path = package_dir / "episode-manifest.json"
    candidates = [source_manifest]
    if manifest_path.exists():
        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        source_value = manifest.get("source_manifest")
        if source_value:
            source_path = Path(str(source_value))
            candidates.append(source_path if source_path.is_absolute() else ROOT / source_path)
        report_window = manifest.get("report_window") if isinstance(manifest.get("report_window"), dict) else {}
        until = report_window.get("until") if isinstance(report_window, dict) else None
        if until:
            candidates.insert(0, Path("__inline_until__") / str(until))

    for candidate in candidates:
        if str(candidate).startswith("__inline_until__/"):
            until = str(candidate).replace("__inline_until__/", "", 1)
        elif candidate.exists():
            manifest = json.loads(candidate.read_text(encoding="utf-8"))
            until = manifest.get("until")
        else:
            continue
        if not until:
            continue
        try:
            dt = datetime.fromisoformat(str(until).replace("Z", "+00:00"))
            return dt.astimezone(timezone(timedelta(hours=8))).strftime("%y%m%d")
        except ValueError:
            return str(until)[:10].replace("-", "")[2:]
    return None


def episode_count(markdown: str) -> int:
    return len(re.findall(r"^####\s+情报\s+\d+", markdown, flags=re.MULTILINE))


def delivery_name(markdown_path: Path) -> str:
    run_id = markdown_path.stem.replace("-gemini-report", "")
    scheduled_date = scheduled_report_date(run_id)
    if scheduled_date:
        return f"{scheduled_date}-Spotify播客情报研报"
    manifest_path = ROOT / "data" / "gemini_inputs" / run_id / "episode-manifest.json"
    if manifest_path.exists():
        import json

        manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
        episodes = manifest.get("episodes") or []
        if episodes and episodes[0].get("published_at"):
            published_at = episodes[0]["published_at"]
            try:
                published_dt = datetime.fromisoformat(published_at.replace("Z", "+00:00"))
                local_dt = published_dt.astimezone(timezone(timedelta(hours=8)))
                date_part = local_dt.strftime("%y%m%d")
            except ValueError:
                date_part = published_at[:10].replace("-", "")[2:]
            return f"{date_part}-Spotify播客情报研报"
    return f"{short_date(run_id)}-Spotify播客情报研报"


def split_report(markdown: str) -> tuple[str, list[str], str | None, list[Section]]:
    lines = [line.rstrip() for line in markdown.splitlines()]
    title = "Spotify 播客情报研报"
    body_start = 0
    for index, line in enumerate(lines):
        if line.startswith("# "):
            title = line[2:].strip()
            body_start = index + 1
            break

    episodes: list[Section] = []
    intro: list[str] = []
    current: Section | None = None
    second_part_heading: str | None = None
    for raw in lines[body_start:]:
        if raw.startswith("<!--"):
            continue
        if raw.strip().startswith("**Run ID:"):
            continue
        if raw.startswith("## 第二部分"):
            second_part_heading = raw.replace("##", "").strip()
            continue
        if raw.startswith("#### 情报 "):
            if current:
                episodes.append(current)
            current = Section(title=raw.replace("####", "").strip(), lines=[])
            continue
        if current is None:
            intro.append(raw)
        else:
            current.lines.append(raw)
    if current:
        episodes.append(current)
    return title, intro, second_part_heading, episodes


TITLE_OVERRIDES = {
    "20260525-201553": (
        "AI时代下的商业、组织与个人变革情报研报：SaaS韧性、数字员工与人机协作新范式 "
        "(Podcast Intelligence Report: SaaS Resilience, Digital Workers & the New Human-AI Operating Model)"
    ),
    "20260523-222300": (
        "全球科技与宏观前沿情报研报：AI基建、算力主权与后真相时代的文化反击 "
        "(Global Tech & Macro Intelligence Report: AI Infrastructure, Compute Sovereignty & Cultural Countermoves)"
    ),
    "20260603-131000-260601-combined": (
        "AI范式转移与产业重构情报研报：代币短缺、视频代理、算力内存与资本新秩序 "
        "(Podcast Intelligence Report: AI Paradigm Shift, Video Agents, Compute Memory & the New Capital Order)"
    ),
}


def report_display_title(run_id: str, markdown_title: str) -> str:
    return TITLE_OVERRIDES.get(run_id, markdown_title)


PART_TITLE_OVERRIDES = {
    "第一部分": "第一部分：本期核心判断 (Core Judgment)",
    "第二部分": "第二部分：逐集情报与证据 (Episode Intelligence & Evidence)",
    "第三部分": "第三部分：跨节目专题分析 (Cross-Episode Thematic Analysis)",
    "第四部分": "第四部分：第二层思维 (Second-Level Thinking)",
    "第五部分": "第五部分：结论与战略意义 (Conclusion & Strategic Implications)",
}


def normalize_part_title(title: str) -> str:
    for prefix, normalized in PART_TITLE_OVERRIDES.items():
        if title.startswith(prefix):
            return normalized
    return title


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = text.replace("`", "")
    return text.strip()


def should_keep_inline_bold(text: str) -> bool:
    normalized = clean_inline(text).strip()
    if not normalized:
        return False
    generic_markers = (
        "影响",
        "议题",
        "重塑",
        "演进",
        "变革",
        "趋势",
        "方面",
        "现象",
        "经验",
        "安全性",
        "重要性",
    )
    if any(marker in normalized for marker in generic_markers) and len(normalized) > 10:
        return False
    professional_markers = (
        "AI-First",
        "AI代理",
        "SaaS",
        "token",
        "Agent",
        "代理",
        "数字员工",
        "人类三明治",
        "共享团队代理",
        "Harness",
        "AEO",
        "OpenRouter",
        "Cursor",
        "Anthropic",
        "Copilot",
        "Retardmaxxing",
        "Cialis",
        "他达拉非",
        "SaaS末日论",
        "so what",
        "前线部署工程师",
        "品格",
        "character",
    )
    if any(marker in normalized for marker in professional_markers):
        return len(normalized) <= 22
    if re.search(r"[A-Za-z][A-Za-z0-9+/#-]{2,}", normalized):
        return len(normalized) <= 22
    return False


def inline_runs(text: str, italic: bool = False, *, allow_leading_subtitle_bold: bool = False) -> list[TextRun]:
    text = text.strip()
    if allow_leading_subtitle_bold:
        match = re.match(r"^(\s*(?:\d+\.\s*)?)\*\*(.+?)\*\*([：:]?)(.*)$", text)
        if match:
            prefix, title, punctuation, rest = match.groups()
            runs: list[TextRun] = []
            if prefix:
                runs.append(TextRun(prefix, italic=italic))
            runs.append(TextRun(title, bold=True, italic=italic))
            if punctuation or rest:
                runs.append(TextRun(punctuation + rest.replace("*", ""), italic=italic))
            return runs
    parts = re.split(r"(\*\*.*?\*\*)", text)
    runs: list[TextRun] = []
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            content = part[2:-2]
            runs.append(TextRun(content, italic=italic))
        else:
            runs.append(TextRun(part.replace("*", ""), italic=italic))
    return runs


def label_bold_runs(text: str, italic: bool = False, *, allow_leading_subtitle_bold: bool = False) -> list[TextRun]:
    if text.strip().startswith("**"):
        return inline_runs(text, italic=italic, allow_leading_subtitle_bold=allow_leading_subtitle_bold)
    if re.match(r"^\[?\d{0,2}:?\d{1,2}:\d{2}\]?", text.strip()):
        return [TextRun(text.replace("*", ""), italic=italic)]
    if "：" in text:
        label, rest = text.split("：", 1)
        clean_label = clean_inline(label)
        force_label = clean_label.startswith(COMPACT_LABELS)
        if force_label:
            return [TextRun(clean_label, bold=True, italic=italic), TextRun("：" + rest.replace("*", ""), italic=italic)]
    if ":" in text:
        label, rest = text.split(":", 1)
        clean_label = clean_inline(label)
        force_label = clean_label.startswith(COMPACT_LABELS)
        if force_label:
            return [TextRun(clean_label, bold=True, italic=italic), TextRun(":" + rest.replace("*", ""), italic=italic)]
    return inline_runs(text, italic=italic, allow_leading_subtitle_bold=allow_leading_subtitle_bold)


def paragraph_role(line: str) -> tuple[str, str]:
    stripped = line.strip()
    if stripped == "<!-- pagebreak -->":
        return "pagebreak", ""
    if not stripped or stripped == "---":
        return "blank", ""
    if stripped.startswith("## "):
        return "h2", normalize_part_title(clean_inline(stripped[3:]))
    if stripped.startswith("### "):
        return "h3", clean_inline(stripped[4:])
    if stripped.startswith("#### "):
        return "h3", clean_inline(stripped[5:])
    if stripped.startswith("- "):
        return "body", stripped[2:].strip()
    if stripped.startswith("* "):
        return "body", stripped[2:].strip()
    if re.match(r"^\d+\.\s+", stripped):
        return "body", stripped
    return "body", stripped


def set_doc_columns(section, count: int) -> None:
    sect_pr = section._sectPr
    cols = sect_pr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    cols_el.set(qn("w:num"), str(count))
    cols_el.set(qn("w:space"), "420")
    if not cols:
        sect_pr.append(cols_el)


def enable_east_asian_line_break_rules(doc: Document) -> None:
    settings = doc.settings._element
    for tag in (
        "w:kinsoku",
        "w:noLineBreaksBefore",
        "w:noLineBreaksAfter",
        "w:doNotUseEastAsianBreakRules",
        "w:overflowPunct",
    ):
        for element in list(settings.findall(qn(tag))):
            settings.remove(element)

    kinsoku = OxmlElement("w:kinsoku")
    kinsoku.set(qn("w:val"), "1")
    settings.append(kinsoku)

    no_before = OxmlElement("w:noLineBreaksBefore")
    no_before.set(qn("w:val"), "，。；：？！、,.!?;:)）]】》”’％%")
    settings.append(no_before)

    no_after = OxmlElement("w:noLineBreaksAfter")
    no_after.set(qn("w:val"), "（([【《“‘")
    settings.append(no_after)

    settings.append(OxmlElement("w:overflowPunct"))


def add_page_number_footer(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char_end)


def set_run_font(run, font_name: str = DOCX_FONT) -> None:
    run.font.name = font_name
    r_pr = run._element.get_or_add_rPr()
    rfonts = r_pr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        r_pr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), font_name)
    lang = r_pr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        r_pr.append(lang)
    lang.set(qn("w:val"), "zh-CN")
    lang.set(qn("w:eastAsia"), "zh-CN")


def force_bold(run) -> None:
    run.bold = True
    set_run_font(run, DOCX_BOLD_FONT)
    r_pr = run._element.get_or_add_rPr()
    for tag in ("w:b", "w:bCs"):
        element = r_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            r_pr.append(element)
        element.set(qn("w:val"), "1")


def force_style_bold(style) -> None:
    style.font.bold = True
    style.font.name = DOCX_BOLD_FONT
    r_pr = style._element.get_or_add_rPr()
    rfonts = r_pr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        r_pr.append(rfonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{attr}"), DOCX_BOLD_FONT)
    for tag in ("w:b", "w:bCs"):
        element = r_pr.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            r_pr.append(element)
        element.set(qn("w:val"), "1")


def add_horizontal_rule(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BFBFBF")
    p_bdr.append(bottom)


def add_blank_line(doc: Document, points: float = 7.5) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(points)


def paragraph_text_kind(text: str) -> str:
    normalized = clean_inline(text)
    if normalized.startswith(METADATA_LABELS):
        return "metadata"
    if normalized.startswith(CONTENT_BLOCK_LABELS):
        return "content_label"
    return "body"


def apply_body_spacing(paragraph, text: str, previous_kind: str | None) -> str:
    kind = paragraph_text_kind(text)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    paragraph.paragraph_format.space_before = Pt(0)
    if kind == "metadata":
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
    elif kind == "content_label":
        paragraph.paragraph_format.space_before = Pt(7.5 if previous_kind not in {None, "metadata"} else 3)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
    else:
        paragraph.paragraph_format.space_after = Pt(7.5)
    return kind


def apply_reference_paragraph_format(paragraph, role: str, is_first_second_heading: bool = False) -> None:
    paragraph.paragraph_format.line_spacing = 1.15
    if role == "h1":
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(29)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
    elif role == "h2":
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(22)
        paragraph.paragraph_format.space_before = Pt(6 if is_first_second_heading else 0)
        paragraph.paragraph_format.space_after = Pt(3)
    elif role == "h3":
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(17)
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
    elif role == "body":
        paragraph.paragraph_format.space_after = Pt(7.5)


def add_docx_paragraph(doc: Document, role: str, text: str, *, italic: bool = False, allow_leading_subtitle_bold: bool = False):
    if role == "pagebreak":
        doc.add_page_break()
        return None
    if role == "blank":
        return None
    style_names = {style.name for style in doc.styles if style.name}
    style = {
        "h1": "Heading 1",
        "h2": "Heading 2",
        "h3": "Heading 3",
    }.get(role, "Body Text" if "Body Text" in style_names else "normal")
    paragraph = doc.add_paragraph(style=style)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if role in {"h1", "h2", "h3"} else None
    apply_reference_paragraph_format(paragraph, role)
    if role in {"h1", "h2", "h3"}:
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
    leading_subtitle = re.match(r"^\s*(?:\d+\.\s*)?\*\*.+?\*\*([：:]?)(.*)$", text.strip())
    if role == "body" and allow_leading_subtitle_bold and leading_subtitle and len(leading_subtitle.group(2).strip()) <= 24:
        paragraph.paragraph_format.keep_with_next = True
        paragraph.paragraph_format.keep_together = True
    runs = (
        label_bold_runs(text, italic=italic, allow_leading_subtitle_bold=allow_leading_subtitle_bold)
        if role in {"body", "bullet", "number"}
        else [TextRun(text, bold=role in {"h1", "h2", "h3"})]
    )
    for text_run in runs:
        run = paragraph.add_run(text_run.text)
        if text_run.bold:
            force_bold(run)
        if text_run.italic:
            run.italic = True
        if not text_run.bold:
            set_run_font(run)
    return paragraph


def template_document() -> Document:
    doc = Document(REFERENCE_DOCX) if REFERENCE_DOCX.exists() else Document()
    body = doc._body._element
    sect_pr = body.sectPr
    for child in list(body):
        if child is sect_pr:
            continue
        body.remove(child)
    return doc


def report_body_lines(markdown: str) -> Iterable[str]:
    saw_title = False
    for raw in markdown.splitlines():
        line = raw.rstrip()
        if line.strip() == "<!-- pagebreak -->":
            yield line
            continue
        if line.startswith("<!--"):
            continue
        if line.strip().startswith("**Run ID:"):
            continue
        if not saw_title and line.startswith("# "):
            saw_title = True
            continue
        normalized_line = clean_inline(line)
        if normalized_line.startswith("- "):
            normalized_line = normalized_line[2:].strip()
        if normalized_line.startswith("Transcript 来源"):
            continue
        yield line


def is_translation_line(line: str) -> bool:
    stripped = line.strip()
    if not stripped.startswith("*"):
        return False
    if re.search(r"\[\d{0,2}:?\d{1,2}:\d{2}\]|\d{1,2}:\d{2}", stripped):
        return False
    if "Speaker" in stripped or "发言者" in stripped:
        return False
    return True


def add_report_body(doc: Document, markdown: str) -> None:
    seen_part = False
    seen_episode = False
    in_key_quote = False
    current_part_number = 0
    last_role: str | None = None
    previous_body_kind: str | None = None
    for line in report_body_lines(markdown):
        role, text = paragraph_role(line)
        if role == "blank":
            continue

        if role == "h2":
            part_match = re.match(r"^第([一二三四五])部分", text)
            if part_match:
                current_part_number = {"一": 1, "二": 2, "三": 3, "四": 4, "五": 5}[part_match.group(1)]
            if seen_part:
                add_blank_line(doc)
                add_horizontal_rule(doc)
                add_blank_line(doc)
            seen_part = True
            seen_episode = False
            in_key_quote = False
            previous_body_kind = None
        elif role == "h3" and text.startswith("情报 "):
            if seen_episode and last_role != "h2":
                add_blank_line(doc)
                add_horizontal_rule(doc)
                add_blank_line(doc)
            seen_episode = True
            in_key_quote = False
            previous_body_kind = None
        elif role in {"h3", "h1"}:
            in_key_quote = False
            previous_body_kind = None
            if last_role not in {None, "h2"}:
                add_blank_line(doc)

        if role == "body":
            normalized = clean_inline(text)
            if "关键金句" in normalized and "结论" in normalized:
                in_key_quote = True
            elif "证据锚点" in normalized or line.strip() == "---":
                in_key_quote = False

        paragraph = add_docx_paragraph(
            doc,
            role,
            text,
            italic=in_key_quote and is_translation_line(line),
            allow_leading_subtitle_bold=role == "body" and current_part_number >= 3,
        )
        if paragraph is not None and role == "body":
            normalized = clean_inline(text)
            previous_body_kind = apply_body_spacing(paragraph, normalized, previous_body_kind)
        last_role = role


def build_docx(markdown_path: Path, out_path: Path) -> None:
    run_id = markdown_path.stem.replace("-gemini-report", "")
    markdown = markdown_path.read_text(encoding="utf-8")
    title, _, _, _ = split_report(markdown)

    doc = template_document()
    enable_east_asian_line_break_rules(doc)
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(27.94)
    section.page_height = Cm(21.59)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)
    set_doc_columns(section, 1)
    add_page_number_footer(section)

    styles = doc.styles
    for name in ("Normal", "normal", "Body Text", "List Bullet", "List Number"):
        if name not in [style.name for style in styles]:
            continue
        style = styles[name]
        style.font.name = DOCX_FONT
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{attr}"), DOCX_FONT)
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(7.5)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(BODY_LINE_SPACING_PT)
    for name, size, color in (
        ("Heading 1", 24, RGBColor(0, 0, 0)),
        ("Heading 2", 18, RGBColor(0, 0, 0)),
        ("Heading 3", 14, RGBColor(0, 0, 0)),
    ):
        style = styles[name]
        r_pr = style._element.get_or_add_rPr()
        rfonts = r_pr.rFonts
        if rfonts is None:
            rfonts = OxmlElement("w:rFonts")
            r_pr.append(rfonts)
        for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
            rfonts.set(qn(f"w:{attr}"), DOCX_BOLD_FONT)
        style.font.size = Pt(size)
        force_style_bold(style)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12 if name != "Heading 2" else 11.25)
        style.paragraph_format.space_after = Pt(12 if name != "Heading 2" else 11.25)

    p = doc.add_paragraph(style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    apply_reference_paragraph_format(p, "h1")
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.keep_together = True
    r = p.add_run(report_display_title(run_id, title))
    force_bold(r)
    add_report_body(doc, markdown)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)


def register_pdf_font(font_path: str) -> str:
    font_name = "ArialUnicode"
    pdfmetrics.registerFont(TTFont(font_name, font_path))
    return font_name


def pdf_styles(font_name: str) -> dict[str, ParagraphStyle]:
    base = getSampleStyleSheet()
    return {
        "title": ParagraphStyle(
            "DeliveryTitle",
            parent=base["Title"],
            fontName=font_name,
            fontSize=24,
            leading=28,
            alignment=TA_LEFT,
            textColor=colors.black,
            spaceAfter=10,
            wordWrap="CJK",
        ),
        "meta": ParagraphStyle(
            "DeliveryMeta",
            parent=base["Normal"],
            fontName=font_name,
            fontSize=8.6,
            leading=10.4,
            alignment=TA_CENTER,
            textColor=colors.HexColor("#666666"),
            spaceAfter=8,
            wordWrap="CJK",
        ),
        "h1": ParagraphStyle(
            "DeliveryH1",
            parent=base["Heading1"],
            fontName=font_name,
            fontSize=18,
            leading=21,
            textColor=colors.black,
            spaceBefore=3,
            spaceAfter=5,
            keepWithNext=1,
            wordWrap="CJK",
        ),
        "h2": ParagraphStyle(
            "DeliveryH2",
            parent=base["Heading2"],
            fontName=font_name,
            fontSize=18,
            leading=21,
            textColor=colors.black,
            spaceBefore=3,
            spaceAfter=4,
            keepWithNext=1,
            wordWrap="CJK",
        ),
        "h3": ParagraphStyle(
            "DeliveryH3",
            parent=base["Heading3"],
            fontName=font_name,
            fontSize=14,
            leading=16.5,
            textColor=colors.black,
            spaceBefore=3,
            spaceAfter=4,
            keepWithNext=1,
            wordWrap="CJK",
        ),
        "body": ParagraphStyle(
            "DeliveryBody",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=8.5,
            leading=10.1,
            alignment=TA_JUSTIFY,
            spaceAfter=2.8,
            wordWrap="CJK",
        ),
        "bullet": ParagraphStyle(
            "DeliveryBullet",
            parent=base["BodyText"],
            fontName=font_name,
            fontSize=8.5,
            leading=10.1,
            leftIndent=12,
            firstLineIndent=-7,
            alignment=TA_LEFT,
            spaceAfter=2.2,
            wordWrap="CJK",
        ),
    }


def draw_page(canvas, doc) -> None:
    canvas.saveState()
    canvas.setFillColor(colors.white)
    canvas.rect(0, 0, landscape(letter)[0], landscape(letter)[1], stroke=0, fill=1)
    canvas.setFont("ArialUnicode", 7.5)
    canvas.setFillColor(colors.HexColor("#777777"))
    canvas.drawCentredString(landscape(letter)[0] / 2, 0.5 * cm, str(canvas.getPageNumber()))
    canvas.restoreState()


def pdf_paragraph(style_map: dict[str, ParagraphStyle], role: str, text: str) -> Paragraph | Spacer:
    if role == "blank":
        return Spacer(1, 4)
    if role == "body":
        return Paragraph(pdf_inline(label_bold_runs(text)), style_map["body"])
    paragraph = Paragraph(pdf_escape(text), style_map.get(role, style_map["body"]))
    if role in {"h1", "h2", "h3"}:
        paragraph.keepWithNext = 1
    return paragraph


def grouped_pdf_flowables(flowables: list[tuple[str, object]]) -> list[object]:
    grouped: list[object] = []
    index = 0
    heading_roles = {"h1", "h2", "h3"}
    while index < len(flowables):
        role, flowable = flowables[index]
        if role not in heading_roles:
            grouped.append(flowable)
            index += 1
            continue

        block = [flowable]
        index += 1
        while index < len(flowables) and flowables[index][0] in heading_roles:
            block.append(flowables[index][1])
            index += 1
        if index < len(flowables):
            block.append(flowables[index][1])
            index += 1
        grouped.append(CondPageBreak(6.2 * cm))
        grouped.append(KeepTogether(block))
    return grouped


def pdf_escape(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def pdf_inline(runs: list[TextRun]) -> str:
    formatted: list[str] = []
    for run in runs:
        text = pdf_escape(run.text)
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        formatted.append(text)
    return "".join(formatted)


def build_pdf(markdown_path: Path, out_path: Path, font_path: str) -> None:
    run_id = markdown_path.stem.replace("-gemini-report", "")
    markdown = markdown_path.read_text(encoding="utf-8")
    markdown_title, intro, second_part_heading, episodes = split_report(markdown)
    font_name = register_pdf_font(font_path)
    styles = pdf_styles(font_name)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc = SimpleDocTemplate(
        str(out_path),
        pagesize=landscape(letter),
        rightMargin=2.54 * cm,
        leftMargin=2.54 * cm,
        topMargin=2.54 * cm,
        bottomMargin=2.54 * cm,
    )
    story: list = [Paragraph(report_display_title(run_id, markdown_title), styles["title"])]
    flowables: list[tuple[str, object]] = []
    for line in intro:
        role, text = paragraph_role(line)
        if role == "blank":
            continue
        flowables.append((role, pdf_paragraph(styles, role, text)))
    for index, episode in enumerate(episodes):
        if index == 0 and second_part_heading:
            heading = Paragraph(second_part_heading, styles["h2"])
            heading.keepWithNext = 1
            flowables.append(("h2", heading))
        heading = Paragraph(episode.title, styles["h3"])
        heading.keepWithNext = 1
        flowables.append(("h3", heading))
        for line in episode.lines:
            role, text = paragraph_role(line)
            if role == "blank":
                continue
            flowables.append((role, pdf_paragraph(styles, role, text)))
    story.extend(grouped_pdf_flowables(flowables))
    doc.build(story, onFirstPage=draw_page, onLaterPages=draw_page)


def export_docx_to_pdf_with_word(docx_path: Path, pdf_path: Path) -> bool:
    if pdf_path.exists():
        pdf_path.unlink()
    script = f'''
set inputPath to POSIX file "{docx_path}"
set outputPath to POSIX file "{pdf_path}"
tell application "Microsoft Word"
    open inputPath
    delay 1
    set reportDocument to active document
    save as reportDocument file name outputPath file format format PDF
    close reportDocument saving no
end tell
'''
    completed = subprocess.run(["osascript"], input=script, text=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
    if completed.returncode != 0 or not pdf_path.exists():
        details = "\n".join(
            part
            for part in (
                f"returncode={completed.returncode}",
                f"stdout={completed.stdout.strip()}",
                f"stderr={completed.stderr.strip()}",
            )
            if part and not part.endswith("=")
        )
        print(f"Word PDF export failed: {details}", file=sys.stderr)
    return completed.returncode == 0 and pdf_path.exists()


def render(markdown_path: Path, font_path: str, *, allow_reportlab_fallback: bool = False) -> tuple[Path, Path]:
    run_id = markdown_path.stem.replace("-gemini-report", "")
    stem = delivery_name(markdown_path)
    docx_path = ROOT / "reports" / "word" / f"{stem}.docx"
    pdf_path = ROOT / "reports" / "pdf" / f"{stem}.pdf"
    build_docx(markdown_path, docx_path)
    if not export_docx_to_pdf_with_word(docx_path, pdf_path):
        if not allow_reportlab_fallback:
            raise RuntimeError(
                f"Microsoft Word PDF export failed for {docx_path}. "
                "The ReportLab fallback does not match the reference Word formatting; "
                "rerun with Word available or pass --allow-reportlab-fallback explicitly."
            )
        build_pdf(markdown_path, pdf_path, font_path)
    return docx_path, pdf_path


def main() -> None:
    parser = argparse.ArgumentParser(description="Render reviewed Markdown reports into delivery DOCX/PDF files.")
    parser.add_argument("reports", nargs="+", type=Path)
    parser.add_argument("--font", default=DEFAULT_FONT)
    parser.add_argument(
        "--allow-reportlab-fallback",
        action="store_true",
        help="Allow the simpler ReportLab PDF fallback if Microsoft Word export is unavailable.",
    )
    args = parser.parse_args()
    for report in args.reports:
        docx_path, pdf_path = render(report, args.font, allow_reportlab_fallback=args.allow_reportlab_fallback)
        print(f"docx={docx_path}")
        print(f"pdf={pdf_path}")


if __name__ == "__main__":
    main()
